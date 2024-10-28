from library.reader import Reader
from library.book import Book
from library.subscription import FreeSubscription, PaidSubscription
from fastapi import FastAPI, HTTPException
from fastapi.responses import FileResponse
from pydantic import BaseModel
import sqlite3
from docx import Document
from openpyxl import Workbook
import uvicorn
import os

app = FastAPI()

DATABASE_PATH = os.getenv("DATABASE_PATH", "/app/data/library.db")
MIGRATIONS_PATH = "/app/migrations.sql"
OUT_DIR = "/app/reports"
os.makedirs(OUT_DIR, exist_ok=True)

class ReaderModel(BaseModel):
    name: str

class BookModel(BaseModel):
    title: str
    author: str
    pages: int
    reader_id: int

class StatisticsRequestModel(BaseModel):
    reader_id: int
    subscription_type: str
    max_pages: int = 0

def get_db_connection():
    """Функция для создания подключения к базе данных."""
    return sqlite3.connect(DATABASE_PATH)

def initialize_database():
    """Инициализация базы данных с выполнением миграций из migrations.sql."""
    if os.path.exists(MIGRATIONS_PATH):
        with get_db_connection() as conn:
            cursor = conn.cursor()
            with open(MIGRATIONS_PATH, 'r') as f:
                sql_script = f.read()
            cursor.executescript(sql_script)
            conn.commit()

initialize_database()

@app.post("/add_reader/")
async def add_reader(reader: ReaderModel):
    with get_db_connection() as conn:
        cursor = conn.cursor()
        cursor.execute("INSERT INTO readers (name) VALUES (?)", (reader.name,))
        conn.commit()
        reader_id = cursor.lastrowid
    return {"message": f"Reader {reader.name} added", "id": reader_id}

@app.post("/add_book/")
async def add_book(book: BookModel):
    with get_db_connection() as conn:
        cursor = conn.cursor()
        cursor.execute("INSERT INTO books (title, author, pages, reader_id) VALUES (?, ?, ?, ?)", 
                       (book.title, book.author, book.pages, book.reader_id))
        conn.commit()
    return {"message": f"Book {book.title} by {book.author} added"}


@app.post("/calculate_statistics/")
async def calculate_statistics(request: StatisticsRequestModel):
    with get_db_connection() as conn:
        cursor = conn.cursor()
        
        # Получаем данные читателя
        cursor.execute("SELECT name FROM readers WHERE id = ?", (request.reader_id,))
        reader_data = cursor.fetchone()
        if not reader_data:
            raise HTTPException(status_code=404, detail="Reader not found")

        reader = Reader(name=reader_data[0], reader_id=request.reader_id)

        # Получаем книги, прочитанные конкретным читателем
        cursor.execute("SELECT title, author, pages FROM books WHERE reader_id = ?", (request.reader_id,))
        books = cursor.fetchall()
        for book_data in books:
            book = Book(title=book_data[0], author=book_data[1], pages=book_data[2])
            reader.read_book(book)

        # Определяем тип подписки
        if request.subscription_type == "Free":
            subscription = FreeSubscription(reader, request.max_pages)
        else:
            subscription = PaidSubscription(reader, request.max_pages)

        stats = subscription.calculate_statistics()

        # Сохраняем подписку
        cursor.execute("INSERT INTO subscriptions (reader_id, type, max_pages) VALUES (?, ?, ?)", 
                       (request.reader_id, request.subscription_type, request.max_pages))
        conn.commit()

    return stats

@app.get("/generate_report/")
async def generate_report(format: str = "docx"):
    file_path = f"{OUT_DIR}/report.{format}"

    if format == "docx":
        doc = Document()
        doc.add_heading("Library Report", 0)
        with get_db_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT * FROM readers")
            for reader in cursor.fetchall():
                doc.add_paragraph(f"Reader ID: {reader[0]}, Name: {reader[1]}")
        doc.save(file_path)
        mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif format == "xlsx":
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "Library Report"
        with get_db_connection() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT * FROM readers")
            for idx, reader in enumerate(cursor.fetchall(), 1):
                sheet.append([reader[0], reader[1]])
        workbook.save(file_path)
        mime_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    else:
        raise HTTPException(status_code=400, detail="Invalid format specified")
    
    # Возвращаем файл с заданным MIME-типом
    return FileResponse(path=file_path, filename=f"report.{format}", media_type=mime_type)

@app.get("/check_db/")
async def check_db():
    with get_db_connection() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT name FROM sqlite_master WHERE type='table';")
        tables = cursor.fetchall()
    return {"tables": tables}

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8081)
